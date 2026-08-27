#!/usr/bin/env python3
"""Generate a Word-compatible weekly progress report."""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).with_name(
    "Pulse-and-Collector-Progress-and-Target-Report-2026-08-11-to-08-22.docx"
)


def set_run_font(run, *, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=18, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    else:
        set_run_font(run, size=13, bold=True, color=RGBColor(0x2E, 0x75, 0xB6))
    return p


def add_meta(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=11, bold=True)
    return p


def add_item(doc, title, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    title_run = p.add_run(title)
    set_run_font(title_run, size=11, bold=True)
    body_run = p.add_run(" " + body)
    set_run_font(body_run, size=11)
    return p


def add_numbered(doc, n, text, bold_bits=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{n}. {text}")
    set_run_font(run, size=11)
    return p


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    add_heading(doc, "Weekly Progress Report & Targets", 1)

    add_meta(doc, "Progress Period: August 11 - August 18, 2026")
    add_meta(doc, "Targets Period: August 19 - August 22, 2026")
    add_meta(doc, "Apps: Pulse (payroll / timekeeping) and Biometric Collector")

    add_heading(doc, "I. Pulse - August 11 - August 16, 2026", 2)

    items = [
        (
            "Master File upload - multi-campus slots:",
            "Template/upload now supports campus assignments 1-5 (primary required; slots 2-5 optional).",
        ),
        (
            "Master File upload cleanup:",
            "Removed unused JSON collection columns from the master-file template and import.",
        ),
        (
            "Master File upsert:",
            'Added "Disable required fields" so partial rows can import; Employee Number + Email match updates an existing employee, otherwise creates.',
        ),
        (
            "Employee Credentials tab:",
            "Per-employee file list with Add modal, preview, download, soft delete, and audit logs; later linked to Document Types.",
        ),
        (
            "Credential preview:",
            "In-app preview for images, PDF, Word, and Excel; LibreOffice PDF conversion when available (desktop-safe).",
        ),
        (
            "Document Types (HR):",
            "Maintenance lookup with Required flag; seeded onboarding checklist; government IDs split (TIN, SSS, PhilHealth, Pag-IBIG, PRC).",
        ),
        (
            "Employee Loans tab:",
            "PATHS-style loans on Employee Profile after Credentials (modal CRUD, payment schemes, sys_logs).",
        ),
        (
            "Payroll Shift Code (per day):",
            "Add Shift Code on batch employee detail plus Upload Adjustments tab; Process/Reprocess applies the override.",
        ),
        (
            "Manual Overtime:",
            "Add Overtime on batch detail (date + start/end vs excess hours); bulk Overtime CSV upload; policy flags ignored for manual filings.",
        ),
        (
            "Attendance day details:",
            "Minutes on late/undertime deductions; clickable Late / Undertime / Overtime rows show work dates.",
        ),
        (
            "Last punch as Time Out:",
            "If the last log of the day is tagged In, payroll still uses it as the session Out.",
        ),
        (
            "Offset absent tardiness with OT:",
            "Policy option to keep regular hours and reduce OT when late would have marked the day absent.",
        ),
        (
            "Blank Grace Period:",
            "Timekeeping Policy grace can be left empty (no grace).",
        ),
        (
            "Employee Profile Calendar View:",
            "Month calendar of first In / last Out; day click lists all punches.",
        ),
        (
            "PATHS-style Attendance View:",
            "Daily summary (shift, in/out, basic, excess, OT, ND, late, undertime); date-from / date-to; PDF; Reports module for multiple employees.",
        ),
        (
            "Attendance View vs payroll:",
            "OT/late follow processed batch rules; later saved as day snapshots on Process.",
        ),
        (
            "S3 Pull logs:",
            "Time Logs can pull gzipped collector JSON from S3 (biometric_logs/) and attach punches by campus biometric ID.",
        ),
        (
            "Payroll Register:",
            "Campus worksheets (skip empty campuses); Staff/Admin employee-type layout.",
        ),
        (
            "BIR 1601-C / 2316:",
            "Official ENCS blanks; MWE items 15-16; optional whole-year 13th month; multi-select posted batches in the same month.",
        ),
        (
            "Desktop offline:",
            "App runs without internet; S3 update/backup checks skip when offline.",
        ),
        (
            "Desktop builds:",
            "Pulse 0.1.50 through 0.1.58 (paired macOS + Windows). Latest: v0.1.58 on S3 payroll_installer/.",
        ),
    ]
    for title, body in items:
        add_item(doc, title, body)

    add_heading(doc, "I. Pulse - August 16, 2026 (wrap-up)", 2)
    add_item(
        doc,
        "Attendance day snapshots:",
        "Process / Reprocess writes trn_payroll_attendance_days so Attendance View and Late/UT/OT modals read saved payroll metrics instead of recomputing every time. Legacy batches still fall back until reprocessed.",
    )
    add_item(
        doc,
        "Desktop build 0.1.58:",
        "Shipped the snapshot feature in paired installers.",
    )

    add_heading(doc, "II. Biometric Collector - August 13 and August 17, 2026", 2)
    collector_items = [
        (
            "Installer download Server Error:",
            'Windows "Update required" no longer proxies the ~150MB EXE through PHP; uses an S3 pre-signed URL. Shipped v0.1.5.',
        ),
        (
            "Collect now Server Error (San Mateo):",
            "First pull of a large K30 buffer (~99k punches) ran out of memory in NativePHP. Collect now now runs in a background worker; chunked SQLite inserts and S3 gzip uploads; JSON error messages instead of a generic 500. Shipped v0.1.6.",
        ),
        (
            "Attendance start date:",
            "Dashboard start date skips punches older than the window after device fetch (before insert/export); SQLite-safe S3 update chunks. Shipped v0.1.7.",
        ),
        (
            "Loading overlay:",
            "Full-screen loader on page loads, navigation, forms, and Collect now; hidden when done. Background status polls do not flash the overlay.",
        ),
        (
            "Log retention (Months to keep):",
            "Keep only the last N months in SQLite. Older logs are archived as local JSON under storage/app/biometric-deleted-logs/, then deleted. Collect now and auto-collect are blocked until months is set. Shipped v0.1.8.",
        ),
        (
            "Save retention 500 fix:",
            "Saving months only stores the setting (no prune in the HTTP request). JSON archive + delete runs on the next collect. Missing column is created on boot/save. Shipped v0.1.9.",
        ),
        (
            "Desktop builds:",
            "Collector 0.1.5 through 0.1.9. Latest: v0.1.9 on S3 biometric_installer/.",
        ),
    ]
    for title, body in collector_items:
        add_item(doc, title, body)

    add_heading(doc, "Targets - August 19 - August 22, 2026", 2)

    add_heading(doc, "I. Pulse", 2)
    add_item(
        doc,
        "UAT of 0.1.58:",
        "Attendance View vs processed batch, Staff/Admin Payroll Register, BIR 1601-C multi-batch / 13th month, and S3 Pull logs.",
    )
    add_item(
        doc,
        "S3 Pull logs E2E:",
        "Confirm collector uploads land in biometric_logs/, Pulse Pull logs creates Time Logs batches, and Process applies punches.",
    )
    add_item(
        doc,
        "Web S3 documents to desktop viewing:",
        "Pull employee documents that were uploaded to S3 through the web, then provide viewing on Pulse desktop (Credentials / Document Types: list, preview, download; reuse existing PDF/Office preview where possible).",
    )
    add_item(
        doc,
        "Report / payroll fixes from UAT:",
        "Apply feedback on Attendance View, Staff register, and BIR stamps; ship a follow-up Pulse build only after UAT sign-off.",
    )

    add_heading(doc, "II. Biometric Collector", 2)
    add_item(
        doc,
        "San Mateo update to 0.1.9:",
        "Install latest, set Months to keep, then Collect now.",
    )
    add_item(
        doc,
        "Confirm collect + S3:",
        "Logs under biometric_logs/2026/08/San-Mateo/ and Pulse can pull them.",
    )
    add_item(
        doc,
        "Watch first-collect duration:",
        "Device still downloads its full attendance buffer; start date only filters after pull. If still too slow, evaluate clearing device logs after a successful collect.",
    )
    add_item(
        doc,
        "Stability:",
        "No new collector features until San Mateo collect and Pulse import are verified.",
    )

    add_heading(doc, "III. Cross-app (end-to-end)", 2)
    add_numbered(doc, 1, "Collector 0.1.9 collect to S3.")
    add_numbered(doc, 2, "Pulse Pull logs, then Process.")
    add_numbered(doc, 3, "Web S3 employee documents to Pulse desktop viewing.")

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
