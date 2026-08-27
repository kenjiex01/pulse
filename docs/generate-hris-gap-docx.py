#!/usr/bin/env python3
"""Generate a user-friendly People360 vs HRIS Blueprint status Word document.

Includes Desktop (pulse/) and Web (skolaris-fe/be /people360) columns.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).with_name("People360-HRIS-Blueprint-Status-2026-08-27.docx")

NAVY = RGBColor(0x0B, 0x31, 0x8F)
CYAN = RGBColor(0x00, 0xA3, 0xE6)
GRAY = RGBColor(0x47, 0x55, 0x69)
AMBER = RGBColor(0x92, 0x40, 0x0E)
RED = RGBColor(0x99, 0x1B, 0x1B)


def set_run_font(run, *, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=20, bold=True, color=NAVY)
    elif level == 2:
        set_run_font(run, size=14, bold=True, color=NAVY)
    else:
        set_run_font(run, size=12, bold=True, color=CYAN)
    return p


def add_para(doc, text, *, size=11, bold=False, color=GRAY, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=11, color=GRAY)
    return p


def tint_status(cell, status: str):
    status = str(status)
    if status == "Partial":
        shade_cell(cell, "FEF3C7")
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=8, bold=True, color=AMBER)
    elif status in ("Not yet built", "Pending"):
        shade_cell(cell, "FEE2E2")
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=8, bold=True, color=RED)
    elif status == "Ready":
        shade_cell(cell, "DCFCE7")


def add_table(doc, headers, rows, header_fill="0B318F", status_cols=None):
    """status_cols: set of header names to tint by cell value."""
    status_cols = status_cols or {"Status", "Desktop", "Web", "Combined"}
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        set_run_font(run, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
        shade_cell(hdr[i], header_fill)
    for r_i, row in enumerate(rows):
        cells = table.rows[r_i + 1].cells
        for c_i, value in enumerate(row):
            cells[c_i].text = ""
            p = cells[c_i].paragraphs[0]
            run = p.add_run(str(value))
            set_run_font(run, size=8, bold=(c_i == 0))
            if headers[c_i] in status_cols:
                tint_status(cells[c_i], value)
    doc.add_paragraph()
    return table


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    add_heading(doc, "People360", 1)
    p = doc.add_paragraph()
    run = p.add_run("HRIS Blueprint Status Report")
    set_run_font(run, size=16, bold=True, color=CYAN)
    add_para(
        doc,
        "Desktop + Web · What already works · What is incomplete · What is not yet built",
        size=11,
        color=RGBColor(0x64, 0x74, 0x8B),
    )
    add_para(doc, "Date: 27 August 2026  |  Audience: Management, HR, and project stakeholders", size=10)
    add_para(
        doc,
        "This report compares the Philippine HRIS Software Module Blueprint (28 modules) with People360 Desktop (pulse app) and People360 Web (Skolaris /people360 in iskolaris-fe + iskolaris-be). Combined = best of both. Planning guide only — not a legal review.",
        size=10,
    )

    add_heading(doc, "How to read this report", 2)
    add_para(doc, "Three statuses, shown for Desktop, Web, and Combined:", size=11)
    add_bullet(doc, "Ready — Fully matches the blueprint for day-to-day use. (None of the 28 are 100% ready yet.)")
    add_bullet(doc, "Partial — Useful screens/workflows exist, but important blueprint items are still missing.")
    add_bullet(doc, "Not yet built — No matching module or workflow on that product.")
    add_bullet(doc, "Desktop = People360 installer app (pulse/). Web = browser People360 under /people360 (iskolaris-fe + iskolaris-be).")

    add_heading(doc, "Snapshot — Combined (Desktop + Web)", 2)
    add_para(
        doc,
        "Desktop is strongest for school payroll and timekeeping. Web is strongest for employee self-service, document review, attendance checker, and Job Order requests. Together they are still not a full hire-to-retire HRIS.",
    )
    add_table(
        doc,
        ["Status", "Count", "Meaning"],
        [
            ["Ready", "0 of 28", "No module is complete vs the full blueprint."],
            ["Partial", "13 of 28", "Includes M16 (self-service) now that Web ESS exists."],
            ["Not yet built", "15 of 28", "Needs new design and development on both products."],
        ],
    )
    add_para(doc, "In plain words:", bold=True, color=NAVY)
    add_bullet(doc, "Desktop: manage employees, run payroll, pull time logs, generate BIR / SSS / PhilHealth / Pag-IBIG reports, faculty load pay.")
    add_bullet(doc, "Web: employees log in, update profile, upload documents for HR review, view attendance, file Job Order (WRF) requests, and get approvals.")
    add_bullet(doc, "Still missing on both: recruitment, leave filing/balances, contractors, discipline cases, full clearance/separation, live clock-in.")

    add_heading(doc, "What already works well", 2)
    add_heading(doc, "On Desktop", 3)
    for title, body in [
        ("Payroll", "Pay periods, compute/post, OT / ND / holiday, payslips, register, BIR 1601-C / 2316 / Alphalist."),
        ("Government contributions", "SSS, PhilHealth, Pag-IBIG tables and contribution reports; employee loans."),
        ("Timekeeping", "Policies, shifts, holidays, time log upload, biometric pull, attendance edit, OT for payroll."),
        ("Employee records", "Faculty / staff / admin / hybrid profiles; campuses; salary history; credentials."),
        ("Faculty load", "Pull from Skolaris, upload loads, pay faculty from load hours."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(title)
        set_run_font(r1, size=11, bold=True, color=NAVY)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(6)
        r2 = p2.add_run(body)
        set_run_font(r2, size=10, color=GRAY)

    add_heading(doc, "On Web (Skolaris People360)", 3)
    for title, body in [
        ("Employee self-service (M16)", "Login/OTP, dashboard, profile edit, own attendance & loading attendance, My Requests, request approvals."),
        ("Documents (M04)", "Employee upload / mark N/A; HR document review (verify, remap type, approve/reject)."),
        ("Attendance ops (M05)", "Timekeeping (load-based), Attendance Checker (Present/Absent/Late, offline), biometric log restore."),
        ("Local employees (M02)", "Browse/edit restored desktop employees; sync APIs for the desktop app."),
        ("Faculty load PDFs (M20)", "Upload/parse/manage faculty loading PDFs in People360."),
        ("Privacy & access (M15 / M01)", "Privacy consent gate, Pulse audit trail, users/roles, desktop API keys."),
        ("Backup restore", "Payroll SQL backups and restored data tables (ops — not live payroll)."),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r1 = p.add_run(title)
        set_run_font(r1, size=11, bold=True, color=NAVY)
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(6)
        r2 = p2.add_run(body)
        set_run_font(r2, size=10, color=GRAY)

    add_heading(doc, "Full checklist — Desktop · Web · Combined", 2)
    all_rows = [
        ["M01", "Organization & HR Configuration", "MVP", "Partial", "Partial", "Partial"],
        ["M02", "Worker / Employee Master Data", "MVP", "Partial", "Partial", "Partial"],
        ["M03", "Recruitment & Applicant Tracking", "MVP", "Not yet built", "Not yet built", "Not yet built"],
        ["M04", "Onboarding & Employment Documents", "MVP", "Partial", "Partial", "Partial"],
        ["M05", "Time, Attendance & Scheduling", "MVP", "Partial", "Partial", "Partial"],
        ["M06", "Leave & Absence Management", "MVP", "Partial", "Not yet built", "Partial"],
        ["M07", "Payroll & Statutory Pay", "MVP", "Partial", "Not yet built", "Partial"],
        ["M08", "Benefits & Government Contributions", "MVP", "Partial", "Not yet built", "Partial"],
        ["M09", "Performance Management", "Phase 2", "Not yet built", "Not yet built", "Not yet built"],
        ["M10", "Learning & Development", "Phase 2", "Not yet built", "Not yet built", "Not yet built"],
        ["M11", "Grievance & Discipline", "MVP", "Not yet built", "Not yet built", "Not yet built"],
        ["M12", "Separation, Clearance & Retirement", "MVP", "Not yet built", "Not yet built", "Not yet built"],
        ["M13", "Contractor & Outsourced Workforce", "MVP", "Not yet built", "Not yet built", "Not yet built"],
        ["M14", "Occupational Safety & Health", "MVP", "Not yet built", "Not yet built", "Not yet built"],
        ["M15", "Data Privacy & Records", "MVP", "Partial", "Partial", "Partial"],
        ["M16", "Self-Service, Requests & Approvals", "MVP", "Not yet built", "Partial", "Partial"],
        ["M17", "HR Compliance & Legal Rules", "MVP", "Partial", "Not yet built", "Partial"],
        ["M18", "HR Analytics, Dashboards & Audit", "Phase 2", "Partial", "Partial", "Partial"],
        ["M19", "Faculty & Academic Personnel", "MVP (schools)", "Partial", "Partial", "Partial"],
        ["M20", "Faculty Load, Schedule & Overload", "MVP (schools)", "Partial", "Partial", "Partial"],
        ["M21", "Faculty Evaluation & Promotion", "Phase 2", "Not yet built", "Not yet built", "Not yet built"],
        ["M22", "Faculty Development & Research", "Phase 2", "Not yet built", "Not yet built", "Not yet built"],
        ["M23", "Student-Safeguarding / Conduct", "Phase 2", "Not yet built", "Not yet built", "Not yet built"],
        ["M24", "Sales / Commission & Incentive", "Optional", "Not yet built", "Not yet built", "Not yet built"],
        ["M25", "Shift, Field & Branch Workforce", "Optional", "Not yet built", "Not yet built", "Not yet built"],
        ["M26", "Succession & Workforce Planning", "Phase 3", "Not yet built", "Not yet built", "Not yet built"],
        ["M27", "AI Assistance & Automation", "Phase 3", "Not yet built", "Not yet built", "Not yet built"],
        ["M28", "SDG / ESG Workforce Reporting", "Phase 3", "Not yet built", "Not yet built", "Not yet built"],
    ]
    add_table(
        doc,
        ["ID", "Module", "Priority", "Desktop", "Web", "Combined"],
        all_rows,
    )

    add_heading(doc, "What is partial — short notes", 2)
    add_table(
        doc,
        ["ID", "Area", "Desktop", "Web", "What you should know"],
        [
            ["M01", "Organization", "Partial", "Partial", "Desktop has org masters. Web has users/roles + API keys only."],
            ["M02", "Employee master", "Partial", "Partial", "Desktop = full CRUD. Web = restored local employees + ESS profile."],
            ["M04", "Documents", "Partial", "Partial", "Web ESS upload + HR review. No full onboarding case / e-contract."],
            ["M05", "Time & attendance", "Partial", "Partial", "Desktop import/payroll. Web checker + ESS read. No live clock-in."],
            ["M06", "Leave", "Partial", "Not yet built", "Desktop leave types for payroll only. No filing on either product."],
            ["M07", "Payroll", "Partial", "Not yet built", "Strong on desktop. Web = SQL backup restore only."],
            ["M08", "Benefits", "Partial", "Not yet built", "Statutory tables/reports on desktop only."],
            ["M15", "Privacy", "Partial", "Partial", "Desktop sys_logs. Web privacy consent + Pulse audit."],
            ["M16", "Self-service", "Not yet built", "Partial", "Web ESS + Job Order/WRF approvals. Leave ESS still missing."],
            ["M17", "Legal rules", "Partial", "Not yet built", "Rate/govt tables on desktop; not versioned."],
            ["M18", "Analytics", "Partial", "Partial", "Operational reports + checker analytics. No KPI warehouse."],
            ["M19", "Faculty personnel", "Partial", "Partial", "Faculty/staff/admin model. No PRC / qualified subjects."],
            ["M20", "Faculty load", "Partial", "Partial", "Desktop pay from load. Web PDF upload + loading attendance."],
        ],
    )

    add_heading(doc, "What is not yet built (either product)", 2)
    add_heading(doc, "Needed for a complete MVP (recommended soon)", 3)
    add_table(
        doc,
        ["ID", "Module", "Status", "Why it matters"],
        [
            ["M03", "Recruitment", "Not yet built", "Job requests, posting, applicants, interviews, offers."],
            ["M11", "Discipline & grievance", "Not yet built", "NTE, hearings, sanctions, appeals."],
            ["M12", "Separation & clearance", "Not yet built", "Resignation/termination, clearance, COE, final pay."],
            ["M13", "Contractors / outsourced", "Not yet built", "Agency accreditation, deployed workers, compliance."],
            ["M14", "Occupational Safety & Health", "Not yet built", "Incidents, PPE, drills, corrective actions."],
            ["M06+", "Leave filing (gap)", "Not yet built on Web", "Web ESS exists but leave request/balance is still missing."],
        ],
    )
    add_heading(doc, "Later phases (can wait)", 3)
    add_bullet(doc, "Phase 2 — Performance, L&D, richer analytics, faculty evaluation/promotion, research, student-safeguarding.")
    add_bullet(doc, "Phase 3 — Succession, AI assistance, SDG/ESG reporting.")
    add_bullet(doc, "Optional (private business) — Sales commissions; field/branch workforce ops.")

    add_heading(doc, "Suggested next steps", 2)
    add_para(doc, "From the Combined position (Desktop payroll + Web ESS), close the biggest MVP gaps first:")
    steps = [
        "Leave filing & balances (M06) — use Web ESS shell; Desktop already has leave types.",
        "Finish self-service (M16) — add OT corrections, certificates, leave on top of Job Order.",
        "Effective-dated government rates (M17 / M08) — Desktop, so old payroll stays correct.",
        "Onboarding checklist (M04) — extend Web soft gate + Desktop credentials.",
        "Separation & clearance (M12) — resignation, clearance, final pay, COE.",
        "Recruitment (M03) — if hire-to-retire is in scope.",
        "Complete faculty gaps (M19 / M20) — licenses, qualified subjects, overload approval.",
        "Risk modules (M11 / M13 / M14) — discipline, contractors, safety.",
    ]
    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{i}. {step}")
        set_run_font(run, size=11, color=GRAY)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.paragraph_format.space_before = Pt(18)
    r = footer.add_run("— End of report —")
    set_run_font(r, size=10, italic=True, color=RGBColor(0x94, 0xA3, 0xB8))

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = note.add_run(
        "Source: HRIS Blueprint (PH). Products: People360 Desktop (pulse/) + Web (iskolaris-fe/be /people360). "
        "Technical detail: pulse/docs/hris-blueprint-gap-analysis.md"
    )
    set_run_font(r2, size=8, color=RGBColor(0x94, 0xA3, 0xB8))

    OUT.parent.mkdir(parents=True, exist_ok=True)
    if OUT.exists():
        OUT.unlink()
    doc.save(str(OUT))
    print(f"Wrote {OUT}")
    print(f"Size: {OUT.stat().st_size} bytes")


if __name__ == "__main__":
    main()
