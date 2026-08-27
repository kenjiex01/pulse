#!/usr/bin/env python3
"""Generate Word weekly progress report for Aug 17-23 progress / Aug 24-28 targets."""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).with_name(
    "Pulse-and-Collector-Progress-and-Target-Report-2026-08-17-to-08-28.docx"
)


def set_run_font(run, *, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 12)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    if level == 1:
        set_run_font(run, size=18, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    else:
        set_run_font(run, size=13, bold=True, color=RGBColor(0x2E, 0x75, 0xB6))
    return p


def add_meta(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    set_run_font(run, size=11, bold=True)
    return p


def add_item(doc, title, body):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.15
    title_run = p.add_run(title)
    set_run_font(title_run, size=11, bold=True)
    body_run = p.add_run(" " + body)
    set_run_font(body_run, size=11)
    return p


def add_numbered(doc, n, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(f"{n}. {text}")
    set_run_font(run, size=11)
    return p


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    add_heading(doc, "Weekly Progress Report & Targets", 1)

    add_meta(doc, "Progress Period: August 17 - August 23, 2026")
    add_meta(doc, "Targets Period: August 24 - August 28, 2026")
    add_meta(doc, "Apps: Pulse (payroll / timekeeping) and Biometric Collector")

    add_heading(doc, "Campus rollout (last week)", 2)
    add_item(
        doc,
        "Biometric Collector installed:",
        "Campus PCs for Antipolo, Binangonan, and Taytay. These sites can collect device punches locally, set retention, and upload logs to S3 for Pulse Pull logs. Follow-up this week: confirm Collect now + S3 path + Pulse import for each campus.",
    )

    add_heading(doc, "I. Pulse - August 17 - August 23, 2026", 2)
    pulse_items = [
        (
            "ISKOLARIS Approve (Employee Management):",
            "Pending profile updates show with employee full name; Approve per row, Approve selected (checked only), and Approve all (all shown after search). Does not overwrite salary, credentials, or loans. Matched by Pulse employee ID / employee number.",
        ),
        (
            "S3 Pull logs retry:",
            "Partially imported biometric files are re-read on every pull; already-imported punches stay as duplicates; only new matches are inserted. No empty batch when a retry has nothing new.",
        ),
        (
            "Payroll Register Excel:",
            "One worksheet per campus (empty campuses omitted); names as Last, First Middle, sorted by last name. Uses Main assignment campus; optional Under Campus (e.g. Greenhills under Cainta) rolls employees onto the parent sheet.",
        ),
        (
            "Main assignment + Employee Assignment upload:",
            "Campus cards have a Main assignment checkbox; Employee Assignment upload template can bulk-change the main campus.",
        ),
        (
            "Campus minimum wage:",
            "Nullable minimum wage field on campus maintenance.",
        ),
        (
            "HR Employee report:",
            "Full employee details report (personal, gov IDs, contact, assignments, salary, credentials as Yes/blank); hybrid salary cells tagged Faculty/Staff/Admin when hybrid.",
        ),
        (
            "Encrypted secrets:",
            "Skolaris Pulse API key and S3 credentials can be stored encrypted in .env (enc:...); decrypt only in memory.",
        ),
        (
            "Desktop build:",
            "Pulse 0.1.59 (macOS + Windows) on S3 payroll_installer/ — includes Payroll Register per-campus Excel and last-name format.",
        ),
    ]
    for title, body in pulse_items:
        add_item(doc, title, body)

    add_heading(doc, "II. Biometric Collector - August 17 - August 23, 2026", 2)
    collector_items = [
        (
            "Collect now Server Error (large devices):",
            "Background worker for Collect now; chunked SQLite inserts and S3 gzip uploads; JSON errors instead of generic 500. Build 0.1.6.",
        ),
        (
            "Attendance start date:",
            "Dashboard start date skips punches older than the window after device fetch. Build 0.1.7.",
        ),
        (
            "Loading overlay:",
            "Full-screen loader on navigation, forms, and Collect now (status polls do not flash it).",
        ),
        (
            "Log retention (Months to keep):",
            "Keep only the last N months in SQLite; older logs archived as local JSON then deleted. Collect blocked until months is set. Build 0.1.8.",
        ),
        (
            "Save retention 500 fix:",
            "Saving months only stores the setting; prune runs on next collect. Build 0.1.9 (latest shipped installer on S3 biometric_installer/).",
        ),
        (
            "Deleted log backups:",
            "Nav to list/download retention JSON backups; Upload backup JSON restores kind: deleted_attendance punches (skips duplicates).",
        ),
        (
            "Download DTR:",
            "Date from / date to + multi-select enrolled users to Cainta-style CSV timesheet (Employee: Name (id) with Date / In / Out).",
        ),
        (
            "Edit device:",
            "Dashboard Edit beside Users / Logs / Remove (name, model, IP, port, comm key).",
        ),
        (
            "Encrypt S3 API keys:",
            "DB_BACKUP_S3_KEY / DB_BACKUP_S3_SECRET support enc:... at rest.",
        ),
        (
            "Campus installs:",
            "Antipolo, Binangonan, and Taytay PCs installed with the Collector app.",
        ),
    ]
    for title, body in collector_items:
        add_item(doc, title, body)

    add_heading(doc, "Targets - August 24 - August 28, 2026", 2)

    add_heading(doc, "I. Pulse", 2)
    add_item(
        doc,
        "UAT of 0.1.59:",
        "Payroll Register per-campus sheets + Main assignment / Under Campus; Employee report; ISKOLARIS Approve (per row / selected / all).",
    )
    add_item(
        doc,
        "S3 Pull logs E2E with new campuses:",
        "Confirm Antipolo, Binangonan, and Taytay collector uploads land under biometric_logs/ and Pulse Pull logs + Process apply punches (including retry after employee add).",
    )
    add_item(
        doc,
        "Web S3 documents to desktop viewing:",
        "Pull employee documents uploaded via web S3 and view them on Pulse desktop (Credentials / Document Types: list, preview, download).",
    )
    add_item(
        doc,
        "Report / payroll fixes from UAT:",
        "Apply feedback; ship a follow-up Pulse build only after sign-off.",
    )

    add_heading(doc, "II. Biometric Collector", 2)
    add_item(
        doc,
        "Ship build with Aug 20 features:",
        "Package Download DTR, deleted-log backup download/upload, Edit device, and encrypted S3 secrets into a new installer (next version after 0.1.9) and upload to S3.",
    )
    add_item(
        doc,
        "Verify Antipolo / Binangonan / Taytay:",
        "Set Months to keep, Collect now, confirm S3 objects, then Pulse pull.",
    )
    add_item(
        doc,
        "San Mateo follow-through:",
        "Confirm collect + S3 + Pulse import still stable after 0.1.9 (or newer) update.",
    )
    add_item(
        doc,
        "Rollout next campuses:",
        "Continue installs beyond Antipolo, Binangonan, and Taytay as scheduled.",
    )
    add_item(
        doc,
        "Stability:",
        "Prefer field verification and installer ship over large new features until the three new campuses are green on collect to S3 to Pulse.",
    )

    add_heading(doc, "III. Cross-app (end-to-end)", 2)
    add_numbered(doc, 1, "Collector on Antipolo, Binangonan, Taytay collect to S3.")
    add_numbered(doc, 2, "Pulse Pull logs, then Process for those campus paths.")
    add_numbered(doc, 3, "Ship Collector build that includes DTR + backup modules; update campus PCs as needed.")
    add_numbered(doc, 4, "Web S3 employee documents to Pulse desktop viewing.")

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
